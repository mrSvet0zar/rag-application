"""Document ingestion helpers: text extraction (DOCX/HTML) and URL safety.

Parsing and validation here are pure and synchronous, so they are cheap to
unit-test. The one async helper, `read_capped`, deliberately takes a byte
stream rather than an HTTP response so it can be exercised without a network.
"""

from __future__ import annotations

import io
import ipaddress
import socket
from collections.abc import AsyncIterator
from urllib.parse import urlparse

from app.errors import PayloadTooLargeError

# Tags whose text is boilerplate, not content. `math` is in here because
# MathML carries a LaTeX <annotation> twin, so keeping it duplicates every
# formula as unreadable markup.
_STRIP_TAGS = [
    "script",
    "style",
    "noscript",
    "header",
    "footer",
    "nav",
    "aside",
    "form",
    "svg",
    "math",
]

# Elements that stand on their own line in the extracted text. Anything else
# (links, emphasis, spans) is inline and must not break the sentence it sits in.
_BLOCK_TAGS = [
    "p",
    "li",
    "dd",
    "dt",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "blockquote",
    "pre",
    "td",
    "th",
    "caption",
    "figcaption",
]


def validate_public_url(url: str) -> None:
    """Raise ValueError unless `url` is an http(s) URL on a public host.

    Basic SSRF guard: rejects non-http schemes and hosts that resolve to
    loopback / private / link-local / reserved addresses (e.g. localhost,
    10.x, 169.254.169.254 cloud-metadata). Note: this validates the *initial*
    host only — following redirects or DNS rebinding could still reach a
    private address, an accepted limitation for this demo.
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("seuls les schémas http(s) sont autorisés")
    host = parsed.hostname
    if not host:
        raise ValueError("URL sans nom d'hôte")

    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror as exc:
        raise ValueError(f"hôte introuvable ({exc})") from exc

    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_reserved
            or ip.is_multicast
            or ip.is_unspecified
        ):
            raise ValueError("l'URL pointe vers une adresse non publique")


def html_to_text(html: str) -> tuple[str | None, str]:
    """Extract (title, readable_text) from an HTML string.

    Drops scripts/styles/nav/etc., then rebuilds the text one *block* at a
    time. Extracting the document as a whole with a newline separator would
    break a line at every tag boundary, so a sentence containing links comes
    out shredded into one-word lines — which then embeds badly. Joining inside
    a block and separating only between blocks keeps sentences intact.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")

    # Capture the title, then drop the whole <head> so its title/meta/style
    # text doesn't leak into the body text.
    title = None
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    if soup.head:
        soup.head.decompose()

    for tag in soup(_STRIP_TAGS):
        tag.decompose()

    blocks = [
        block
        for block in soup.find_all(_BLOCK_TAGS)
        # Skip containers of other blocks (a <td> wrapping <p>s, say), whose
        # text would otherwise be emitted twice.
        if not block.find(_BLOCK_TAGS)
    ]

    if blocks:
        lines = [block.get_text(" ", strip=True) for block in blocks]
    else:
        # No structure to go on (a bare fragment): fall back to the whole thing.
        lines = [soup.get_text(" ", strip=True)]

    return title, "\n".join(line for line in lines if line)


def docx_to_text(content: bytes) -> str:
    """Extract text from a .docx file's bytes (paragraphs + table cells)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def pdf_to_text(content: bytes) -> str:
    """Extract text from a PDF's bytes."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch to the right extractor based on the file extension.

    Anything unrecognised is treated as plain text (utf-8, then latin-1).
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return pdf_to_text(content)
    if lower.endswith(".docx"):
        return docx_to_text(content)
    if lower.endswith((".html", ".htm")):
        return html_to_text(content.decode("utf-8", errors="replace"))[1]
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


async def read_capped(stream: AsyncIterator[bytes], limit: int, what: str) -> bytes:
    """Accumulate `stream`, aborting as soon as it exceeds `limit` bytes.

    Buffering the whole body first and checking its length afterwards would
    spend the memory before noticing — which is the entire problem this guards
    against.
    """
    buffer = bytearray()
    async for piece in stream:
        buffer.extend(piece)
        if len(buffer) > limit:
            megabytes = limit / (1024 * 1024)
            raise PayloadTooLargeError(
                f"{what} trop volumineux (limite : {megabytes:.0f} Mo)."
            )
    return bytes(buffer)
